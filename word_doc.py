from docx import Document
from re import compile


document = Document('orig_test.docx')

keyword = compile(r'\s*Creator: (?P<creator_name>\w+)\s*')

word_replacement = 'Daniel Serrano'

for _p in document.paragraphs: # doc = docx.Document(...) 

    found_word = keyword.match(_p.text)

    if found_word: 

        for _r in _p.runs: 
            _r.text = _r.text.replace(
                found_word.group('creator_name'), 
                word_replacement
            )

paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')

prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

document.add_heading('The REAL meaning of the universe')

document.add_heading('The role of dolphins', level=2)

document.add_page_break()

table = document.add_table(rows=1, cols=3)

# get table data -------------
items = (
    (7, '1024', 'Plush kittens'),
    (3, '2042', 'Furbees'),
    (1, '1288', 'French Poodle Collars, Deluxe'),
)


# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description'

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]


table.style = 'LightShading-Accent1'

document.save('test.docx')
